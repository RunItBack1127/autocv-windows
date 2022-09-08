from flask import Flask, request, jsonify
from flask_cors import CORS

from docx import Document
from docx.shared import Pt

from github import Github

import os
import uuid
import zipfile

HEADLESS_WORD_MODE = 1
PDF_SAVE_MODE = 17

RESUME_CV_FONT_NAME = "Radian Book"
RESUME_CV_FONT_SIZE = 12

TAB_TO_SPACES = "      "

# Define paths for all application roles
SOFTWARE_ENGINEER_PATH__MICROSERVICES_RESUME = "templates/resume/microservices/WPG_Software_Engineer.docx.zip"
FRONT_END_ENGINEER_PATH__MICROSERVICES_RESUME = "templates/resume/microservices/WPG_Front_End_Engineer.docx.zip"
FULL_STACK_ENGINEER_PATH__MICROSERVICES_RESUME = "templates/resume/microservices/WPG_Full_Stack_Engineer.docx.zip"

SOFTWARE_ENGINEER_PATH__DATABASES_RESUME = "templates/resume/databases/WPG_Software_Engineer.docx.zip"
FRONT_END_ENGINEER_PATH__DATABASES_RESUME = "templates/resume/databases/WPG_Front_End_Engineer.docx.zip"
FULL_STACK_ENGINEER_PATH__DATABASES_RESUME = "templates/resume/databases/WPG_Full_Stack_Engineer.docx.zip"

SOFTWARE_ENGINEER_PATH__MICROSERVICES_GENERATED_RESUME = "templates/resume/microservices/generated/WPG_Software_Engineer.docx.zip"
FRONT_END_ENGINEER_PATH__MICROSERVICES_GENERATED_RESUME = "templates/resume/microservices/generated/WPG_Front_End_Engineer.docx.zip"
FULL_STACK_ENGINEER_PATH__MICROSERVICES_GENERATED_RESUME = "templates/resume/microservices/generated/WPG_Full_Stack_Engineer.docx.zip"

SOFTWARE_ENGINEER_PATH__DATABASES_GENERATED_RESUME = "templates/resume/databases/generated/WPG_Software_Engineer.docx.zip"
FRONT_END_ENGINEER_PATH__DATABASES_GENERATED_RESUME = "templates/resume/databases/generated/WPG_Front_End_Engineer.docx.zip"
FULL_STACK_ENGINEER_PATH__DATABASES_GENERATED_RESUME = "templates/resume/databases/generated/WPG_Full_Stack_Engineer.docx.zip"

SOFTWARE_ENGINEER_PATH__COVER_LETTER = "templates/cover_letter/WPG_Software_Engineer.docx"
FRONT_END_ENGINEER_PATH__COVER_LETTER = "templates/cover_letter/WPG_Front_End_Engineer.docx"
FULL_STACK_ENGINEER_PATH__COVER_LETTER = "templates/cover_letter/WPG_Full_Stack_Engineer.docx"

SOFTWARE_ENGINEER_PATH__GENERATED_COVER_LETTER = "templates/cover_letter/generated/WPG_Software_Engineer.docx"
FRONT_END_ENGINEER_PATH__GENERATED_COVER_LETTER = "templates/cover_letter/generated/WPG_Front_End_Engineer.docx"
FULL_STACK_ENGINEER_PATH__GENERATED_COVER_LETTER = "templates/cover_letter/generated/WPG_Full_Stack_Engineer.docx"

def get_abs_path(filename):
    return os.path.abspath(filename)

# Initialize Flask server
app = Flask(__name__)
CORS(app, resources={r'/*': {'origins': '*'}})

github = Github(os.environ.get("GITHUB_TOKEN"))
repo = github.get_user().get_repo("autocv-cover-letters")

"""
    GET endpoint for populating the resume template,
    with options for modifying the languages and
    toggling the bottom skill between microservices
    and databases

    RETURNS the link to the generated resume PDF
    as retrieved from the GitHub REST API
"""
@app.route('/resume', methods=['GET'])
def generate_resume():
    input_filename = ""
    output_filename = ""

    applicant_role = request.args["applicantRole"]
    competency = request.args["competency"]

    if applicant_role == "Software Engineer":
        if competency == "Microservices":
            input_filename = SOFTWARE_ENGINEER_PATH__MICROSERVICES_RESUME
            output_filename = SOFTWARE_ENGINEER_PATH__MICROSERVICES_GENERATED_RESUME
        elif competency == "Databases":
            input_filename = SOFTWARE_ENGINEER_PATH__DATABASES_RESUME
            output_filename = SOFTWARE_ENGINEER_PATH__DATABASES_GENERATED_RESUME
    elif applicant_role == "Front End Engineer":
        if competency == "Microservices":
            input_filename = FRONT_END_ENGINEER_PATH__MICROSERVICES_RESUME
            output_filename = FRONT_END_ENGINEER_PATH__MICROSERVICES_GENERATED_RESUME
        elif competency == "Databases":
            input_filename = FRONT_END_ENGINEER_PATH__DATABASES_RESUME
            output_filename = FRONT_END_ENGINEER_PATH__DATABASES_GENERATED_RESUME
    elif applicant_role == "Full Stack Engineer":
        if competency == "Microservices":
            input_filename = FULL_STACK_ENGINEER_PATH__MICROSERVICES_RESUME
            output_filename = FULL_STACK_ENGINEER_PATH__MICROSERVICES_GENERATED_RESUME
        elif competency == "Databases":
            input_filename = FULL_STACK_ENGINEER_PATH__DATABASES_RESUME
            output_filename = FULL_STACK_ENGINEER_PATH__DATABASES_GENERATED_RESUME

    with zipfile.ZipFile(input_filename, "r") as input_doc, zipfile.ZipFile(output_filename, "w") as output_doc:
        for input_doc_info in input_doc.infolist():
            with input_doc.open(input_doc_info) as input_doc_file:
                content = input_doc_file.read()
                if input_doc_info.filename == "document.xml":

                    relevantSkills = request.args["relevantSkills"].split(",")
                    for (skill_index, skill) in enumerate(relevantSkills):
                        content = content.replace(bytes(f"{{{{L{skill_index + 1}}}}}", "utf-8"), bytes(skill, "utf-8"))

                    output_doc.writestr(f"word/document.xml", content)
                else:
                    output_doc.writestr(input_doc_info.filename, content)
    
    output_doc_renamed = output_filename[:-4]
    pdf_filename = output_doc_renamed.replace(".docx", ".pdf")
    os.rename(output_filename, output_doc_renamed)

    import comtypes
    import comtypes.client

    comtypes.CoInitialize()

    word = comtypes.client.CreateObject('Word.Application')
    word.Visible = False

    doc = word.Documents.Open(get_abs_path(output_doc_renamed))
    doc.SaveAs(get_abs_path(pdf_filename), FileFormat=PDF_SAVE_MODE)
    doc.Close()

    word.Quit()

    comtypes.CoUninitialize()

    github_pdf_path = f"Weston_P_Greene_Resume_{uuid.uuid4()}.pdf"
    pdf_contents = open(pdf_filename, "rb").read()
    github_response = repo.create_file(github_pdf_path, "Appending generated resume file", pdf_contents, branch="master")

    os.remove(output_doc_renamed)
    os.remove(pdf_filename)

    return jsonify(pdf=github_response['content'].download_url)

"""
    GET endpoint for populating the cover letter
    template with options for modifying the name of
    the hiring manager, the name of the role, and
    the company name

    RETURNS the link to the generated cover letter
    PDF  as retrieved from the GitHub REST API
"""
@app.route('/cv', methods=['GET'])
def generate_cover_letter():
    input_filename = ""
    output_filename = ""
    
    applicant_role = request.args["applicantRole"]
    if applicant_role == "Software Engineer":
        input_filename = SOFTWARE_ENGINEER_PATH__COVER_LETTER
        output_filename = SOFTWARE_ENGINEER_PATH__GENERATED_COVER_LETTER
    elif applicant_role == "Front End Engineer":
        input_filename = FRONT_END_ENGINEER_PATH__COVER_LETTER
        output_filename = FRONT_END_ENGINEER_PATH__GENERATED_COVER_LETTER
    elif applicant_role == "Full Stack Engineer":
        input_filename = FULL_STACK_ENGINEER_PATH__COVER_LETTER
        output_filename = FULL_STACK_ENGINEER_PATH__GENERATED_COVER_LETTER

    input_doc = Document(input_filename)

    norm_style = input_doc.styles['Normal']
    norm_font = norm_style.font

    norm_font.name = "Radian Book"
    norm_font.size = Pt(12)
    
    p__hiring_manager = input_doc.paragraphs[6]
    p__name_of_role = input_doc.paragraphs[7]
    p__default_av_toggle = input_doc.paragraphs[9]
    p__company_name = input_doc.paragraphs[10]

    p__hiring_manager.text = p__hiring_manager.text.replace("{{HIRING_MANAGER}}", request.args["recruiterName"])

    nor_text = p__name_of_role.text.split("{{NAME_OF_ROLE}}")

    p__name_of_role.text = ""
    p__name_of_role.add_run(nor_text[0])
    p__name_of_role.add_run(request.args["nameOfRole"]).bold = True
    p__name_of_role.add_run(nor_text[1])

    cover_letter_content = request.args["coverLetterContent"]
    if cover_letter_content == "Self Driving":
        p__default_av_toggle.text = "Having followed the changing self-driving landscape closely since 2015, joining the research team at the EcoPRT autonomous vehicle lab at NC State was my first step in making my mark on the industry. Assuming the role of the lead web developer further enforced my technical skillset, with the goal of redesigning the web presence for showcasing the work of deploying an autonomous vehicle. Maintaining and organizing all of the assets and leveraging each new feature, along with overseeing the design and technical implementation of the website, demonstrates my ability to contribute and dedicate to a project within the scope of a professional environment."

    p__company_name.text = p__company_name.text.replace("{{COMPANY_NAME}}", request.args["companyName"])

    input_doc.save(output_filename)
    pdf_filename = output_filename.replace(".docx", ".pdf")

    import comtypes
    import comtypes.client

    comtypes.CoInitialize()

    word = comtypes.client.CreateObject('Word.Application')
    word.Visible = False

    doc = word.Documents.Open(get_abs_path(output_filename))
    doc.SaveAs(get_abs_path(pdf_filename), FileFormat=PDF_SAVE_MODE)
    doc.Close()

    word.Quit()

    comtypes.CoUninitialize()

    github_pdf_path = f"Weston_P_Greene_Cover_Letter_{uuid.uuid4()}.pdf"
    pdf_contents = open(pdf_filename, "rb").read()
    github_response = repo.create_file(github_pdf_path, "Appending generated cover letter file", pdf_contents, branch="master")

    os.remove(output_filename)
    os.remove(pdf_filename)

    return jsonify(pdf=github_response['content'].download_url)

"""
    GET endpoint for retrieving a formatted template cover letter
    in raw text for the end user - this text is then persisted to
    the system clipboard for applications that do not have a file
    field for the cover letter
"""
@app.route('/copy', methods=['GET'])
def copy_cover_letter():
    recruiter_name = request.args["recruiterName"]
    company_name = request.args["companyName"]
    name_of_role = request.args["nameOfRole"]
    cover_letter_content = request.args["coverLetterContent"]

    default_or_av_paragraph = ""
    if cover_letter_content == "Self Driving":
        default_or_av_paragraph = "Having followed the changing self-driving landscape closely since 2015, joining the research team at the EcoPRT autonomous vehicle lab at NC State was my first step in making my mark on the industry. Assuming the role of the lead web developer further enforced my technical skillset, with the goal of redesigning the web presence for showcasing the work of deploying an autonomous vehicle. Maintaining and organizing all of the assets and leveraging each new feature, along with overseeing the design and technical implementation of the website, demonstrates my ability to contribute and dedicate to a project within the scope of a professional environment."
    elif cover_letter_content == "Default":
        default_or_av_paragraph = "Contributing in a team on a large project across many different microservices also improved my overall technical skillset, placing me in many different roles for various parts of the development cycle - working in the DevEx team at SailPoint required me to sometimes assume frontend responsibilities, other times required a richer set of backend skills, and a few times would require me to exercise more of my full stack expertise. Ensuring effective communication across teams and within our own team was critical, and deploying a consistent and functional product for the end user was kept as our paramount goal, in spite of a specific set of technologies necessary."

    contents = f"Dear {recruiter_name},\n\n{TAB_TO_SPACES}I wanted to reach out to you to further express my interest in the {name_of_role} position. As a recent graduate of the Computer Science program at North Carolina State University, a Software Engineer Intern for SailPoint Technologies, and as the current lead web developer for the EcoPRT autonomous vehicle lab at NC State, I believe that my technically sound and detailed set of front end skills and coding architecture expertise, along with my ability to collaborate, communicate, and implement punctually within a team, would make me an ideal candidate for this role.\n\n{TAB_TO_SPACES}I possess a strong facility in an array of frontend languages, notably Vue, Node.js and Typescript, along with a robust set of backend technologies in Java, C and C++, and Python. My capacity for developing alone for either personal projects, coursework, or independent lab work would integrate seamlessly with a larger team – I can be expected to fulfill expectations concisely, accurately, and under demanding and changing time constraints without sacrificing quality or project longevity.\n\n{TAB_TO_SPACES}{default_or_av_paragraph}\n\n{TAB_TO_SPACES}My passion and dedication are the qualities I strive to bring to the position and the team. I am confident that I can work diligently with my potential co-workers and further promote the development process of {company_name}. Thank you for your consideration, and I look forward to hearing from you to discuss the position further.\n\n\nSincerely,\nWeston P. Greene"

    return jsonify(contents=contents)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8000)
